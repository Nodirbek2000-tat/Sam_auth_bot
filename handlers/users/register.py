from aiogram import types
from aiogram.dispatcher import FSMContext
import json
import os
import tempfile
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import io

from loader import dp, db, bot
from states.states import RegisterState
from utils.subscription import check_and_request_subscription
from keyboards.inline.buttons import get_options_keyboard, get_confirm_response_keyboard


@dp.callback_query_handler(text="user:register", state='*')
async def start_register(callback: types.CallbackQuery, state: FSMContext):
    await state.finish()

    if not await check_and_request_subscription(bot, db, callback.message):
        await callback.answer()
        return

    profile = await db.get_user_profile(callback.from_user.id)

    if not profile:
        await callback.message.edit_text(
            "⚠️ Avval ma'lumotlaringizni to'ldiring!\n\n"
            "/start - Boshlash"
        )
        await callback.answer()
        return

    if not profile['is_approved']:
        await callback.message.edit_text(
            "⚠️ Sizning profilingiz hali tasdiqlanmagan!\n\n"
            "Admin tasdiqlashini kuting."
        )
        await callback.answer()
        return

    survey = await db.get_active_survey()

    if not survey:
        await callback.message.edit_text(
            "⚠️ Hozirda aktiv so'rovnoma yo'q.\n"
            "Iltimos, keyinroq qayta urinib ko'ring."
        )
        await callback.answer()
        return

    fields = await db.get_survey_fields(survey['id'])

    if not fields:
        await callback.message.edit_text(
            "⚠️ So'rovnomada savollar yo'q.\n"
            "Iltimos, keyinroq qayta urinib ko'ring."
        )
        await callback.answer()
        return

    await state.update_data(
        survey_id=survey['id'],
        survey_name=survey['name'],
        fields=[dict(f) for f in fields],
        current_field=0,
        answers={}
    )

    await send_question(callback.message, state, edit=True)
    await RegisterState.answering.set()
    await callback.answer()


@dp.message_handler(chat_type=types.ChatType.PRIVATE, commands=['register'], state='*')
async def cmd_register(message: types.Message, state: FSMContext):
    await state.finish()

    if not await check_and_request_subscription(bot, db, message):
        return

    profile = await db.get_user_profile(message.from_user.id)

    if not profile:
        await message.answer(
            "⚠️ Avval ma'lumotlaringizni to'ldiring!\n\n"
            "/start - Boshlash"
        )
        return

    if not profile['is_approved']:
        await message.answer(
            "⚠️ Sizning profilingiz hali tasdiqlanmagan!\n\n"
            "Admin tasdiqlashini kuting."
        )
        return

    survey = await db.get_active_survey()

    if not survey:
        await message.answer(
            "⚠️ Hozirda aktiv so'rovnoma yo'q.\n"
            "Iltimos, keyinroq qayta urinib ko'ring."
        )
        return

    fields = await db.get_survey_fields(survey['id'])

    if not fields:
        await message.answer(
            "⚠️ So'rovnomada savollar yo'q.\n"
            "Iltimos, keyinroq qayta urinib ko'ring."
        )
        return

    await state.update_data(
        survey_id=survey['id'],
        survey_name=survey['name'],
        fields=[dict(f) for f in fields],
        current_field=0,
        answers={}
    )

    await send_question(message, state, edit=False)
    await RegisterState.answering.set()


async def send_question(message: types.Message, state: FSMContext, edit: bool = False):
    data = await state.get_data()
    fields = data['fields']
    current = data['current_field']

    if current >= len(fields):
        await show_confirmation(message, state, edit)
        return

    field = fields[current]
    question_num = current + 1
    total = len(fields)

    text = (
        f"📋 <b>{data['survey_name']}</b>\n\n"
        f"❓ Savol {question_num}/{total}:\n\n"
        f"<b>{field['question_text']}</b>"
    )

    if field['field_type'] == 'choice' and field['options']:
        keyboard = get_options_keyboard(field['options'], current)
        if edit:
            await message.edit_text(text, reply_markup=keyboard)
        else:
            await message.answer(text, reply_markup=keyboard)

    elif field['field_type'] == 'photo':
        text += "\n\n📷 Rasm yuboring:"
        if edit:
            await message.edit_text(text)
        else:
            await message.answer(text)

    elif field['field_type'] == 'location':
        text += "\n\n📍 Lokatsiyani yuboring:"
        if edit:
            await message.edit_text(text)
        else:
            await message.answer(text)
    else:
        if edit:
            await message.edit_text(text)
        else:
            await message.answer(text)


async def show_confirmation(message: types.Message, state: FSMContext, edit: bool = False):
    data = await state.get_data()
    fields = data['fields']
    answers = data['answers']

    text = f"📋 <b>{data['survey_name']}</b>\n\n"
    text += "✅ <b>Javoblaringizni tekshiring:</b>\n\n"

    for i, field in enumerate(fields):
        answer = answers.get(str(i))

        if field['field_type'] == 'photo':
            text += f"<b>{field['column_name']}:</b> 📷 Rasm\n"
        elif field['field_type'] == 'location':
            if answer:
                loc = json.loads(answer) if isinstance(answer, str) else answer
                link = f"https://maps.google.com/?q={loc['latitude']},{loc['longitude']}"
                text += f"<b>{field['column_name']}:</b> <a href='{link}'>📍 Xaritada ko'rish</a>\n"
            else:
                text += f"<b>{field['column_name']}:</b> —\n"
        else:
            text += f"<b>{field['column_name']}:</b> {answer if answer else '—'}\n"

    text += "\n❓ Tasdiqlaysizmi?"

    if edit:
        await message.edit_text(text, reply_markup=get_confirm_response_keyboard(), disable_web_page_preview=False)
    else:
        await message.answer(text, reply_markup=get_confirm_response_keyboard(), disable_web_page_preview=False)


@dp.callback_query_handler(lambda c: c.data.startswith("answer:"), state=RegisterState.answering)
async def process_choice_answer(callback: types.CallbackQuery, state: FSMContext):
    _, field_order, option_index = callback.data.split(":")
    field_order = int(field_order)
    option_index = int(option_index)

    data = await state.get_data()
    fields = data['fields']

    field = fields[field_order]
    answer = field['options'][option_index]

    answers = data.get('answers', {})
    answers[str(field_order)] = answer

    await state.update_data(
        answers=answers,
        current_field=field_order + 1
    )

    await send_question(callback.message, state, edit=True)
    await callback.answer()


@dp.message_handler(chat_type=types.ChatType.PRIVATE, state=RegisterState.answering,
                    content_types=types.ContentTypes.TEXT)
async def process_text_answer(message: types.Message, state: FSMContext):
    data = await state.get_data()
    current = data['current_field']
    fields = data['fields']
    field = fields[current]

    if field['field_type'] != 'text':
        await message.answer("⚠️ Iltimos, to'g'ri formatda javob bering!")
        return

    answers = data.get('answers', {})
    answers[str(current)] = message.text

    await state.update_data(
        answers=answers,
        current_field=current + 1
    )

    await send_question(message, state, edit=False)


@dp.message_handler(chat_type=types.ChatType.PRIVATE, state=RegisterState.answering,
                    content_types=types.ContentTypes.PHOTO)
async def process_photo_answer(message: types.Message, state: FSMContext):
    data = await state.get_data()
    current = data['current_field']
    fields = data['fields']
    field = fields[current]

    if field['field_type'] != 'photo':
        await message.answer("⚠️ Iltimos, to'g'ri formatda javob bering!")
        return

    answers = data.get('answers', {})
    answers[str(current)] = message.photo[-1].file_id

    await state.update_data(
        answers=answers,
        current_field=current + 1
    )

    await send_question(message, state, edit=False)


@dp.message_handler(chat_type=types.ChatType.PRIVATE, state=RegisterState.answering,
                    content_types=types.ContentTypes.LOCATION)
async def process_location_answer(message: types.Message, state: FSMContext):
    data = await state.get_data()
    current = data['current_field']
    fields = data['fields']
    field = fields[current]

    if field['field_type'] != 'location':
        await message.answer("⚠️ Iltimos, to'g'ri formatda javob bering!")
        return

    answers = data.get('answers', {})
    location = {
        'latitude': message.location.latitude,
        'longitude': message.location.longitude
    }
    answers[str(current)] = json.dumps(location)

    await state.update_data(
        answers=answers,
        current_field=current + 1
    )

    await send_question(message, state, edit=False)


async def generate_word_document(user_id: int, response_data: dict, fields: list):
    """WORD fayl yaratish - faqat foydalanuvchi kiritgan ma'lumotlar"""

    doc = Document()

    # Sarlavha - So'rovnoma nomi sifatida
    title = doc.add_heading("SO'ROVNOMA NATIJALARI", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0, 0, 0)
    title.paragraph_format.space_after = Pt(16)

    # Foydalanuvchi kiritgan barcha ma'lumotlarni ko'rsatish
    temp_images = []

    for i, field in enumerate(fields):
        column_name = field['column_name']
        answer = response_data.get(column_name, "")

        # Lokatsiya - link sifatida ko'rsatish
        if field['field_type'] == 'location':
            p = doc.add_paragraph()
            run = p.add_run(f"{column_name}: ")
            run.bold = True
            if answer:
                try:
                    loc = json.loads(answer) if isinstance(answer, str) else answer
                    p.add_run(f"{loc['latitude']}, {loc['longitude']}")
                except:
                    p.add_run("—")
            else:
                p.add_run("—")
            p.paragraph_format.space_after = Pt(6)
            continue

        # Savol nomi
        p = doc.add_paragraph()
        run = p.add_run(f"{column_name}: ")
        run.bold = True
        p.paragraph_format.space_after = Pt(4)

        # Rasm
        if field['field_type'] == 'photo' and answer:
            try:
                file = await bot.get_file(answer)
                downloaded_file = await bot.download_file(file.file_path)

                img = Image.open(io.BytesIO(downloaded_file.read()))
                img.thumbnail((500, 500), Image.Resampling.LANCZOS)

                temp_path = os.path.join(tempfile.gettempdir(), f"temp_word_img_{i}.png")
                img.save(temp_path, "PNG")
                temp_images.append(temp_path)

                doc.add_picture(temp_path, width=Inches(4))
                last_p = doc.paragraphs[-1]
                last_p.paragraph_format.space_after = Pt(10)

            except Exception as e:
                p.add_run("📷 Rasm yuklanmadi")
                print(f"Rasm yuklashda xato: {e}")

        # Oddiy javob (text, choice)
        elif field['field_type'] != 'photo':
            p.add_run(str(answer) if answer else "—")

    # Faylni saqlash
    current_date = datetime.now()
    file_path = os.path.join(tempfile.gettempdir(),
                             f"natija_{user_id}_{current_date.strftime('%Y%m%d_%H%M%S')}.docx")
    doc.save(file_path)

    # Temp rasmlarni o'chirish
    for temp_img in temp_images:
        try:
            os.remove(temp_img)
        except:
            pass

    return file_path


@dp.callback_query_handler(text="response:confirm", state=RegisterState.answering)
async def confirm_response(callback: types.CallbackQuery, state: FSMContext):
    data = await state.get_data()

    response_data = {}
    for i, field in enumerate(data['fields']):
        response_data[field['column_name']] = data['answers'].get(str(i), "")

    await db.add_survey_response(
        survey_id=data['survey_id'],
        user_id=callback.from_user.id,
        response_data=response_data
    )

    try:
        await callback.message.edit_text("⏳ Ma'lumotnoma tayyorlanmoqda...")

        word_file = await generate_word_document(
            user_id=callback.from_user.id,
            response_data=response_data,
            fields=data['fields']
        )

        user = callback.from_user
        username_part = f"@{user.username}" if user.username else f"ID: {user.id}"
        caption = (
            f"📋 <b>Yangi so'rovnoma</b>\n\n"
            f"👤 Foydalanuvchi: {user.full_name} ({username_part})\n"
            f"📝 So'rovnoma: <b>{data['survey_name']}</b>"
        )

        # Kanalga yuborish
        with open(word_file, 'rb') as f:
            await bot.send_document("@samfayl", f, caption=caption)

        # Barcha adminlarga yuborish
        admins = await db.get_all_admins()
        for admin in admins:
            try:
                with open(word_file, 'rb') as f:
                    await bot.send_document(admin['telegram_id'], f, caption=caption)
            except Exception:
                pass

        os.remove(word_file)

    except Exception as e:
        print(f"WORD yaratishda xato: {e}")

    await callback.message.edit_text(
        "✅ <b>Rahmat!</b>\n\n"
        "Sizning javoblaringiz muvaffaqiyatli saqlandi! 🎉"
    )

    await state.finish()
    await callback.answer("Muvaffaqiyatli saqlandi!")


@dp.callback_query_handler(text="response:cancel", state=RegisterState.answering)
async def cancel_response(callback: types.CallbackQuery, state: FSMContext):
    await state.finish()

    await callback.message.edit_text(
        "❌ Bekor qilindi.\n\n"
        "Qaytadan boshlash uchun /register buyrug'ini bosing."
    )
    await callback.answer()